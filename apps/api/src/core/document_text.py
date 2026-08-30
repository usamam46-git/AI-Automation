"""
core/document_text.py — turn an uploaded file into embeddable chunks.

Pure functions: no database, no network, no object storage. That is why this
sits in `core/` beside `llm_client.py` rather than inside
`modules/knowledge_base/` — the module convention is exactly five files and has
no slot for a helper, and this is the code in the ingestion path most worth
testing hard, which is far easier when nothing has to be mocked to reach it.

Three jobs, in the order the pipeline needs them:

    content_hash(data)  -> the re-ingestion skip
    extract_text(...)   -> bytes to text, per format
    chunk_text(text)    -> text to embeddable windows

## No OCR, deliberately

`ocr_results` exists in the schema and will invite you to fill it. The 15-day
plan rules it out: scanned-document OCR needs Tesseract or a vision model and is
a multi-day detour that adds nothing a demo shows. The consequence is handled
rather than ignored — a PDF whose pages yield no extractable text raises
`UnextractableDocumentError` instead of producing zero chunks, because a
document that silently indexes to nothing is a knowledge base that quietly
answers "I don't know" forever.
"""

from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass
from functools import lru_cache

import tiktoken

# ---------------------------------------------------------------------------
# Accepted formats
# ---------------------------------------------------------------------------

MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIME_MARKDOWN = "text/markdown"
MIME_PLAIN = "text/plain"

SUPPORTED_MIME_TYPES: frozenset[str] = frozenset({MIME_PDF, MIME_DOCX, MIME_MARKDOWN, MIME_PLAIN})

# Browsers and curl are inconsistent about Markdown in particular, so the
# extension is the tie-breaker when the declared type is generic or absent.
_EXTENSION_MIME: dict[str, str] = {
    ".pdf": MIME_PDF,
    ".docx": MIME_DOCX,
    ".md": MIME_MARKDOWN,
    ".markdown": MIME_MARKDOWN,
    ".txt": MIME_PLAIN,
    ".text": MIME_PLAIN,
}


class DocumentTextError(RuntimeError):
    """Base for every failure in this module."""


class UnsupportedDocumentError(DocumentTextError):
    """The file type is not one we extract. Not retryable."""


class UnextractableDocumentError(DocumentTextError):
    """The type is supported but no text came out — a scan, or a corrupt file."""


def resolve_mime_type(file_name: str, declared: str | None) -> str:
    """
    Decide a file's type from its extension first, its declared type second.

    The extension wins because clients lie by omission far more often than by
    commission: `curl -F file=@policy.md` sends `application/octet-stream`, and
    several browsers send that for `.docx` too. A declared type is only consulted
    when the extension is unknown.
    """
    lowered = file_name.lower()
    for extension, mime in _EXTENSION_MIME.items():
        if lowered.endswith(extension):
            return mime

    if declared:
        base = declared.split(";")[0].strip().lower()
        if base in SUPPORTED_MIME_TYPES:
            return base

    raise UnsupportedDocumentError(
        f"'{file_name}' is not a supported document. Accepted: PDF, DOCX, Markdown, plain text. "
        "Scanned documents are not supported — there is no OCR step."
    )


# ---------------------------------------------------------------------------
# Hashing
# ---------------------------------------------------------------------------


def content_hash(data: bytes) -> str:
    """
    SHA-256 of the raw bytes, hex.

    Drives the re-ingestion skip: re-uploading an unchanged file must not pay to
    embed it again. Over raw bytes rather than extracted text on purpose — it is
    cheaper, and it is computed before extraction so an unchanged file skips the
    parse as well as the embedding.
    """
    return hashlib.sha256(data).hexdigest()


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    page_count: int | None


def _decode_text(data: bytes) -> str:
    """UTF-8, then Latin-1 as a never-fails fallback."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


# A line has to appear on this fraction of pages before it counts as furniture
# rather than content. 0.6 clears a running header on a 5-page policy while
# leaving a sentence that happens to recur on two pages of a long document alone.
_RUNNING_LINE_PAGE_RATIO = 0.6
_DIGIT_RUN = re.compile(r"\d+")

# How many lines at each end of a page can be furniture. A running header or
# footer is by definition at an edge; restricting the search there is what keeps
# digit-normalisation from eating body text (see `_strip_running_lines`).
_RUNNING_LINE_EDGE = 3


def _strip_running_lines(pages: list[str], ratio: float = _RUNNING_LINE_PAGE_RATIO) -> list[str]:
    """
    Drop the running header/footer a paginated document repeats on every page.

    Every real policy PDF carries them — a title bar, a confidentiality mark, a
    page number — and they are pure noise three times over: they are the same on
    every page, so they push every chunk's embedding toward every other chunk's;
    they consume the token budget that should hold policy text; and they are the
    one thing in the document guaranteed to answer no question anyone asks.

    Measured on a real 5-page policy: removing two such lines and splitting on
    headings lifted the gap between the right chunk and the best wrong one from
    +0.031 to +0.085 cosine — the number that decides whether a score floor means
    anything.

    Digits are normalised before counting so `Page 1` and `Page 2` are recognised
    as one recurring line rather than five distinct ones. Comparison is on the
    stripped line, but the ORIGINAL is what gets dropped, so indentation cannot
    make a header survive.

    **A heading is never a candidate, and only the outermost
    `_RUNNING_LINE_EDGE` lines of a page are**,
    and that restriction is doing real work rather than saving time. Digit
    normalisation alone is far too eager in the middle of a page: a document
    whose sections are headed `Article 1`, `Article 2`, ... collapses to one key
    `Article #` that appears on most pages, and stripping those would delete the
    headings — which, since `_mark_headings` now chunks on them, is the worst
    possible line to lose. A running header is at an edge by definition, so
    looking only there costs nothing and removes the failure mode. Caught by
    `test_repeated_body_text_that_differs_only_in_digits_is_not_treated_as_furniture`.

    Deliberately conservative: a single-page document has nothing to compare
    against and is returned untouched, and `ratio` requires a genuine majority.
    Removing a line that turns out to be content is worse than keeping noise.
    """
    if len(pages) < 2:
        return pages

    def key(line: str) -> str:
        return _DIGIT_RUN.sub("#", line.strip())

    def edges(lines: list[str]) -> list[int]:
        """Indices of the lines near the top and bottom of a page."""
        if len(lines) <= 2 * _RUNNING_LINE_EDGE:
            return list(range(len(lines)))
        return [*range(_RUNNING_LINE_EDGE), *range(len(lines) - _RUNNING_LINE_EDGE, len(lines))]

    per_page: list[tuple[list[str], set[int]]] = []
    counts: dict[str, int] = {}
    for page in pages:
        lines = page.splitlines()
        # A heading is never furniture, and it is the single most costly line to
        # lose now that `_mark_headings` chunks on it. `Section 1` / `Section 2`
        # at the top of consecutive pages normalises to one key and would
        # otherwise be stripped by the majority rule even at the page edge.
        candidates = {i for i in edges(lines) if lines[i].strip() and not _HEADING_LINE.match(lines[i].strip())}
        per_page.append((lines, candidates))
        # Per page, not per occurrence: a line repeated twice on one page is not
        # thereby a header.
        for unique in {key(lines[i]) for i in candidates}:
            counts[unique] = counts.get(unique, 0) + 1

    threshold = max(2, int(len(pages) * ratio))
    running = {line for line, count in counts.items() if count >= threshold}
    if not running:
        return pages

    kept = [[line for i, line in enumerate(lines) if i not in candidates or key(line) not in running] for lines, candidates in per_page]

    # Safety valve: furniture is a couple of lines, so a rule that would remove
    # half a page has stopped detecting headers and started deleting the
    # document. Reachable on a short page, where the edge windows overlap and
    # every line is a candidate — a five-line page of prose that varies only by
    # digits was emptied outright before this. Bail entirely rather than
    # partially: the whole point is that a document is better with noise in it
    # than with content missing.
    for original, remaining in zip(pages, kept, strict=True):
        before = sum(1 for line in original.splitlines() if line.strip())
        after = sum(1 for line in remaining if line.strip())
        if before and after * 2 < before:
            return pages

    return ["\n".join(lines) for lines in kept]


def _extract_pdf(data: bytes) -> ExtractedDocument:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except (PdfReadError, OSError, ValueError) as exc:
        raise UnextractableDocumentError(f"Could not read PDF: {exc}") from exc

    # Page breaks are paragraph breaks for chunking purposes: joining with a
    # single newline would weld the last line of one page to the first of the
    # next and hide a real boundary from the paragraph packer.
    return ExtractedDocument(text="\n\n".join(_strip_running_lines(pages)), page_count=len(pages))


def _extract_docx(data: bytes) -> ExtractedDocument:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # python-docx raises a wide range on malformed input
        raise UnextractableDocumentError(f"Could not read DOCX: {exc}") from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]
    # Tables carry a lot of a policy document's actual content — thresholds,
    # approval limits, rates. Dropping them would silently lose the parts most
    # worth retrieving.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    # DOCX has no fixed pagination — it is laid out at render time — so
    # page_count stays None rather than being invented.
    return ExtractedDocument(text="\n\n".join(blocks), page_count=None)


def extract_text(data: bytes, mime_type: str, file_name: str) -> ExtractedDocument:
    """
    Extract text from a supported document, or raise.

    Raises `UnsupportedDocumentError` for a type we do not handle and
    `UnextractableDocumentError` when the type is right but nothing came out —
    the scanned-PDF case. Both are permanent: the ingestion task must not retry
    either.
    """
    if mime_type == MIME_PDF:
        extracted = _extract_pdf(data)
    elif mime_type == MIME_DOCX:
        extracted = _extract_docx(data)
    elif mime_type in {MIME_MARKDOWN, MIME_PLAIN}:
        extracted = ExtractedDocument(text=_decode_text(data), page_count=None)
    else:
        raise UnsupportedDocumentError(f"No extractor for '{mime_type}' ({file_name}).")

    if not extracted.text.strip():
        raise UnextractableDocumentError(
            f"'{file_name}' contains no extractable text. If it is a scanned document, it cannot be "
            "indexed — there is no OCR step in this pipeline."
        )
    return extracted


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

# The encoding the text-embedding-3-* family uses. Counting with anything else
# would put our chunk sizes and the embedding endpoint's view of them out of
# step — tolerable at 500 tokens, not at a context limit.
_ENCODING_NAME = "cl100k_base"

DEFAULT_MAX_TOKENS = 500
DEFAULT_OVERLAP_TOKENS = 60

_PARAGRAPH_SPLIT = re.compile(r"\n\s*\n+")
_WHITESPACE_RUN = re.compile(r"[ \t]+")

# C0 and C1 controls plus DEL, excluding \n and \t which carry structure.
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b-\x1f\x7f-\x9f]")

# `4. Expense Approval & Authorization` or `## Approval`. See `_mark_headings`.
_HEADING_LINE = re.compile(r"(?:#{1,6}\s+\S|\d{1,2}\.\s+[A-Z])")

# A section shorter than this is merged forward instead of becoming its own
# chunk. A 20-token heading stub embeds to almost nothing useful and would
# outrank real content on a keyword-ish query purely by being short.
MIN_SECTION_TOKENS = 110


@lru_cache(maxsize=1)
def _encoding() -> tiktoken.Encoding:
    """Cached — building an encoding reads and parses a vocabulary file."""
    return tiktoken.get_encoding(_ENCODING_NAME)


def count_tokens(text: str) -> int:
    return len(_encoding().encode(text))


@dataclass(frozen=True)
class TextChunk:
    index: int
    content: str
    token_count: int


def _normalise(text: str) -> str:
    """
    Collapse horizontal whitespace and trim lines; keep blank-line structure.

    Control characters are flattened to a space first. PDF glyph tables routinely
    map bullets and en-dashes onto unmapped code points, and pypdf hands those
    back verbatim — the Afaqhims policy extracts its bullets and its title-bar
    dashes as U+007F. They are replaced with a SPACE rather than a guessed `-`
    or `•`: the byte says only "no glyph mapping", and inventing punctuation from
    it would put characters in the corpus that are not in the document. Collapsing
    to whitespace loses nothing an embedding uses, and the run collapse below
    tidies up after it.
    """
    cleaned = _CONTROL_CHARS.sub(" ", text.replace("\r\n", "\n"))
    lines = [_WHITESPACE_RUN.sub(" ", line).strip() for line in cleaned.split("\n")]
    return "\n".join(lines)


def _mark_headings(text: str) -> str:
    """
    Put a blank line before every heading, so the paragraph packer sees sections.

    Necessary because `pypdf` emits line breaks but almost never blank ones: a
    whole page arrives as a single paragraph, so `_PARAGRAPH_SPLIT` saw five
    units for a five-page policy and the packer welded them into four ~455-token
    chunks covering three or four unrelated sections each. An embedding of that
    is an average of everything in it, which is why the scores compressed into a
    0.28-0.57 band with almost nothing separating a right answer from a wrong one.

    Recognises a numbered heading (`4. Expense Approval`) and a Markdown one
    (`## Approval`). The numbered form requires a capitalised word after the dot,
    which keeps `1. 250,000` and a decimal in prose from being read as headings;
    it is a heuristic, and a false positive costs one extra chunk boundary rather
    than any loss of text.
    """
    out: list[str] = []
    for line in text.split("\n"):
        if out and _HEADING_LINE.match(line.strip()):
            out.append("")
        out.append(line)
    return "\n".join(out)


def _split_oversized(unit: str, max_tokens: int) -> list[str]:
    """
    Break one over-long paragraph on token boundaries.

    Reached by a paragraph that exceeds the window on its own — a wall-of-text
    contract clause, or a PDF page extracted without paragraph breaks. Splitting
    on tokens loses the semantic boundary, which is why it is the fallback and
    not the strategy.
    """
    encoding = _encoding()
    tokens = encoding.encode(unit)
    return [encoding.decode(tokens[start : start + max_tokens]) for start in range(0, len(tokens), max_tokens)]


def chunk_text(
    text: str,
    *,
    max_tokens: int = DEFAULT_MAX_TOKENS,
    overlap_tokens: int = DEFAULT_OVERLAP_TOKENS,
) -> list[TextChunk]:
    """
    Pack text into overlapping, token-bounded chunks on paragraph boundaries.

    Paragraph-first rather than a sliding token window: a window that cuts
    mid-sentence embeds a fragment whose meaning is not the meaning of either
    neighbour, and retrieval quality is what this whole phase exists to serve.
    An oversized paragraph falls back to a token split (see `_split_oversized`).

    `overlap_tokens` carries the tail of one chunk into the head of the next so a
    fact spanning a boundary is retrievable from both sides.

    **No chunk is ever empty or whitespace-only.** `LLMClient.embed()` raises on
    those, and it raises for the whole batch — one blank chunk would fail an
    entire document's ingestion.
    """
    if max_tokens < 1:
        raise ValueError(f"max_tokens must be >= 1, got {max_tokens}")
    if not 0 <= overlap_tokens < max_tokens:
        raise ValueError(f"overlap_tokens must be in [0, max_tokens), got {overlap_tokens} with max_tokens={max_tokens}")

    normalised = _mark_headings(_normalise(text))
    # (text, starts_a_section). Only the FIRST fragment of an oversized block
    # inherits the flag: the token-split remainder continues the same section.
    units: list[tuple[str, bool]] = []
    for block in _PARAGRAPH_SPLIT.split(normalised):
        stripped = block.strip()
        if not stripped:
            continue
        heads = bool(_HEADING_LINE.match(stripped.split("\n", 1)[0]))
        if count_tokens(stripped) > max_tokens:
            parts = [part for part in _split_oversized(stripped, max_tokens) if part.strip()]
            units.extend((part, heads and position == 0) for position, part in enumerate(parts))
        else:
            units.append((stripped, heads))

    chunks: list[TextChunk] = []
    current: list[str] = []
    current_tokens = 0

    def flush(*, carry_overlap: bool = True) -> None:
        nonlocal current, current_tokens
        if not current:
            return
        content = "\n\n".join(current).strip()
        if content:
            chunks.append(TextChunk(index=len(chunks), content=content, token_count=count_tokens(content)))
        # Carry the tail forward as overlap. Whole units only — slicing a unit
        # here would reintroduce the mid-sentence cut the packer just avoided.
        #
        # NOT carried across a section boundary. Overlap exists so a fact
        # spanning a boundary stays retrievable from both sides, which is a
        # statement about an arbitrary cut mid-topic. A heading is the opposite:
        # the author declared the topic ends there, and prepending the previous
        # section's tail to a 130-token chunk would put ~45% foreign text into
        # the very embedding this split exists to sharpen.
        carried: list[str] = []
        carried_tokens = 0
        if carry_overlap:
            for unit in reversed(current):
                unit_tokens = count_tokens(unit)
                if carried_tokens + unit_tokens > overlap_tokens:
                    break
                carried.insert(0, unit)
                carried_tokens += unit_tokens
        current = carried
        current_tokens = carried_tokens

    for unit, starts_section in units:
        unit_tokens = count_tokens(unit)
        # A heading ends the previous chunk — but only once that chunk carries
        # enough to stand on its own. Below the floor the section is merged
        # forward, which is what stops a title page or a two-line clause from
        # becoming a chunk that retrieves badly for everything.
        if current and starts_section and current_tokens >= MIN_SECTION_TOKENS:
            flush(carry_overlap=False)
        elif current and current_tokens + unit_tokens > max_tokens:
            flush()
            # A single unit can exceed the window even after a flush when the
            # overlap tail is large; drop the tail rather than overflow.
            if current_tokens + unit_tokens > max_tokens:
                current, current_tokens = [], 0
        current.append(unit)
        current_tokens += unit_tokens

    if current:
        content = "\n\n".join(current).strip()
        if content:
            chunks.append(TextChunk(index=len(chunks), content=content, token_count=count_tokens(content)))

    return chunks
