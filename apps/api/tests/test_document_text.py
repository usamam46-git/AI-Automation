"""
tests/test_document_text.py — extraction, chunking and hashing.

Pure unit tests: no database, no network, no object storage. `core/document_text`
is deliberately free of all three so this file can hammer it, because file-format
edge cases are where the 15-day plan expects the ingestion phase to overrun.
"""

import io

import pytest

from src.core.document_text import (
    DEFAULT_MAX_TOKENS,
    MIME_DOCX,
    MIME_MARKDOWN,
    MIME_PDF,
    MIME_PLAIN,
    MIN_SECTION_TOKENS,
    UnextractableDocumentError,
    UnsupportedDocumentError,
    _strip_running_lines,
    chunk_text,
    content_hash,
    count_tokens,
    extract_text,
    resolve_mime_type,
)

POLICY = """ACCOUNTS PAYABLE POLICY

1. Approval thresholds.
Invoices at or above USD 5,000 require the written approval of the finance
controller before payment is released.

2. Three-way match.
Every supplier invoice must be matched against its purchase order and the
corresponding goods receipt before it is posted to the ledger.
"""


def _make_pdf(lines: list[str]) -> bytes:
    """A minimal, genuinely text-extractable PDF. No fixture files on disk."""

    def esc(s: str) -> str:
        return s.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")

    content = ["BT", "/F1 11 Tf", "56 760 Td", "14 TL"]
    content += [f"({esc(line)}) Tj T*" for line in lines]
    content.append("ET")
    stream = "\n".join(content).encode()

    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] " b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(out.tell())
        out.write(f"{i} 0 obj\n".encode() + body + b"\nendobj\n")
    xref = out.tell()
    out.write(f"xref\n0 {len(objs) + 1}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for off in offsets:
        out.write(f"{off:010d} 00000 n \n".encode())
    out.write(f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())
    return out.getvalue()


# ---------------------------------------------------------------------------
# Type resolution
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("file_name", "declared", "expected"),
    [
        ("policy.pdf", "application/pdf", MIME_PDF),
        # The extension wins over a generic declared type. curl sends
        # application/octet-stream for everything it cannot guess, and several
        # browsers send it for .docx — trusting the header would reject files we
        # can definitely read.
        ("policy.md", "application/octet-stream", MIME_MARKDOWN),
        ("handbook.docx", "application/octet-stream", MIME_DOCX),
        ("notes.txt", None, MIME_PLAIN),
        ("POLICY.PDF", None, MIME_PDF),
    ],
)
def test_resolve_mime_type_prefers_the_extension(file_name, declared, expected):
    assert resolve_mime_type(file_name, declared) == expected


def test_resolve_mime_type_falls_back_to_a_credible_declared_type():
    assert resolve_mime_type("policy", "text/markdown") == MIME_MARKDOWN


@pytest.mark.parametrize("file_name", ["scan.png", "sheet.xlsx", "archive.zip", "noextension"])
def test_unsupported_types_are_rejected_by_name(file_name):
    with pytest.raises(UnsupportedDocumentError):
        resolve_mime_type(file_name, "application/octet-stream")


# ---------------------------------------------------------------------------
# Hashing
# ---------------------------------------------------------------------------


def test_content_hash_is_stable_and_content_sensitive():
    assert content_hash(b"policy") == content_hash(b"policy")
    assert content_hash(b"policy") != content_hash(b"policy ")
    assert len(content_hash(b"policy")) == 64


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------


def test_extracts_a_real_pdf_with_a_page_count():
    extracted = extract_text(_make_pdf(POLICY.splitlines()), MIME_PDF, "policy.pdf")
    assert "ACCOUNTS PAYABLE POLICY" in extracted.text
    assert "USD 5,000" in extracted.text
    assert extracted.page_count == 1


def test_extracts_docx_including_table_cells():
    import docx

    document = docx.Document()
    document.add_paragraph("Expense policy")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Meals"
    table.rows[0].cells[1].text = "USD 40 per day"
    buffer = io.BytesIO()
    document.save(buffer)

    extracted = extract_text(buffer.getvalue(), MIME_DOCX, "expenses.docx")
    assert "Expense policy" in extracted.text
    # Tables carry the thresholds and rates — the parts most worth retrieving.
    # Dropping them would silently lose the answer to most policy questions.
    assert "USD 40 per day" in extracted.text
    # DOCX has no fixed pagination; inventing a number would be worse than None.
    assert extracted.page_count is None


def test_extracts_markdown_and_plain_text():
    assert "Title" in extract_text(b"# Title\n\nBody.", MIME_MARKDOWN, "a.md").text
    assert "Body" in extract_text(b"Body.", MIME_PLAIN, "a.txt").text


def test_undecodable_bytes_do_not_crash_extraction():
    # Latin-1 fallback: a policy exported from a Windows tool is not a reason to
    # fail an upload.
    assert extract_text(b"caf\xe9 receipts", MIME_PLAIN, "a.txt").text


def test_a_document_with_no_extractable_text_raises_rather_than_indexing_nothing():
    """
    The scanned-PDF case, and the reason it must be loud.

    There is no OCR step by design. A scan that extracted to "" would index zero
    chunks, report `indexed`, and leave a knowledge base that answers "I don't
    know" forever with nothing anywhere explaining why.
    """
    with pytest.raises(UnextractableDocumentError):
        extract_text(_make_pdf([]), MIME_PDF, "scan.pdf")
    with pytest.raises(UnextractableDocumentError):
        extract_text(b"   \n\n  ", MIME_PLAIN, "blank.txt")


def test_corrupt_pdf_raises_the_permanent_error():
    with pytest.raises(UnextractableDocumentError):
        extract_text(b"not a pdf at all", MIME_PDF, "broken.pdf")


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def test_short_text_is_one_chunk():
    chunks = chunk_text("A single short paragraph.")
    assert len(chunks) == 1
    assert chunks[0].index == 0
    assert chunks[0].token_count == count_tokens(chunks[0].content)


def test_every_chunk_stays_within_the_window():
    text = "\n\n".join(f"Paragraph {i}. " + ("word " * 40) for i in range(60))
    chunks = chunk_text(text, max_tokens=200, overlap_tokens=25)
    assert len(chunks) > 1
    for chunk in chunks:
        assert chunk.token_count <= 200, f"chunk {chunk.index} is {chunk.token_count} tokens"


def test_indexes_are_sequential_from_zero():
    text = "\n\n".join(f"Paragraph {i}. " + ("word " * 40) for i in range(30))
    chunks = chunk_text(text, max_tokens=150, overlap_tokens=20)
    assert [c.index for c in chunks] == list(range(len(chunks)))


def test_consecutive_chunks_overlap():
    """A fact spanning a boundary has to be retrievable from both sides."""
    text = "\n\n".join(f"Clause {i} states a distinct rule about approvals." for i in range(40))
    chunks = chunk_text(text, max_tokens=120, overlap_tokens=30)
    assert len(chunks) > 1
    tail = chunks[0].content.split("\n\n")[-1]
    assert tail in chunks[1].content


def test_zero_overlap_is_allowed():
    text = "\n\n".join(f"Clause {i}." for i in range(50))
    chunks = chunk_text(text, max_tokens=60, overlap_tokens=0)
    assert len(chunks) > 1
    assert chunks[0].content.split("\n\n")[-1] not in chunks[1].content


def test_a_single_oversized_paragraph_is_split_under_the_cap():
    """No paragraph break to pack on — the token-split fallback has to engage."""
    chunks = chunk_text("word " * 3000, max_tokens=100, overlap_tokens=10)
    assert len(chunks) > 1
    assert all(c.token_count <= 100 for c in chunks)


def test_no_chunk_is_ever_empty():
    """
    `LLMClient.embed()` raises on a whitespace-only input — and it raises for the
    whole batch, so one blank chunk fails an entire document's ingestion.
    """
    messy = "Real content.\n\n\n\n   \n\n\t\n\nMore real content.\n\n   "
    for chunk in chunk_text(messy, max_tokens=50, overlap_tokens=5):
        assert chunk.content.strip(), "an empty chunk would fail the whole embedding batch"


def test_token_count_matches_the_content_it_describes():
    for chunk in chunk_text(POLICY * 5, max_tokens=100, overlap_tokens=10):
        assert chunk.token_count == count_tokens(chunk.content)


def test_whitespace_only_input_produces_no_chunks():
    assert chunk_text("   \n\n  \t ") == []


@pytest.mark.parametrize(
    ("max_tokens", "overlap"),
    [(0, 0), (100, 100), (100, 150), (100, -1)],
)
def test_invalid_window_parameters_are_rejected(max_tokens, overlap):
    with pytest.raises(ValueError):
        chunk_text("text", max_tokens=max_tokens, overlap_tokens=overlap)


def test_default_window_is_sane_for_a_real_policy():
    chunks = chunk_text(POLICY)
    assert len(chunks) >= 1
    assert all(c.token_count <= DEFAULT_MAX_TOKENS for c in chunks)


# ---------------------------------------------------------------------------
# Running headers/footers, and section-aware chunking (2026-08-30)
# ---------------------------------------------------------------------------


def _section(number: int, title: str, words: int = 90) -> str:
    """A heading plus enough body to clear MIN_SECTION_TOKENS on its own."""
    return f"{number}. {title}\n" + ("policy sentence body text " * words)


WORDS = ["alpha", "bravo", "charlie", "delta", "echo"]


def _page(n: int, *, extra: list[str] | None = None) -> str:
    """A realistic page: furniture at both edges, genuinely distinct prose between."""
    body = [
        f"The {WORDS[n]} clause requires original receipts.",
        f"Finance reviews the {WORDS[n]} threshold each year.",
        f"Travel under the {WORDS[n]} rule uses the lowest fare.",
        f"Hospitality billed to {WORDS[n]} needs a documented purpose.",
        f"Advances against {WORDS[n]} reconcile within ten days.",
        f"Records for {WORDS[n]} are retained per the schedule.",
    ]
    return "\n".join(["ACME Policy - Internal Use", *(extra or []), *body, f"Page {n + 1}"])


def test_running_header_and_footer_are_dropped():
    cleaned = _strip_running_lines([_page(n) for n in range(5)])
    assert all("Internal Use" not in page for page in cleaned)
    # `Page 1` and `Page 5` differ only in digits and must be recognised as one
    # recurring line, not five distinct ones.
    assert all("Page" not in page for page in cleaned)
    # The prose between the edges is untouched.
    assert all(f"The {WORDS[n]} clause requires original receipts." in cleaned[n] for n in range(5))


def test_a_heading_is_never_stripped_as_furniture():
    """
    A heading repeated at a page edge would otherwise meet the majority rule,
    and headings are what `_mark_headings` chunks on — the worst line to lose.
    """
    cleaned = _strip_running_lines([_page(n, extra=["1. Overview"]) for n in range(5)])
    assert all("Internal Use" not in page for page in cleaned)
    assert all("1. Overview" in page for page in cleaned)


def test_a_line_on_a_minority_of_pages_is_content_and_survives():
    pages = [_page(n) for n in range(5)]
    pages[0] = "Shared note.\n" + pages[0]
    pages[1] = "Shared note.\n" + pages[1]
    cleaned = _strip_running_lines(pages)
    assert all("Shared note." in page for page in cleaned[:2])


def test_a_single_page_document_is_never_stripped():
    # Nothing to compare against, so every line is content by definition.
    pages = [_page(0)]
    assert _strip_running_lines(pages) == pages


def test_indentation_cannot_smuggle_a_header_through():
    cleaned = _strip_running_lines([_page(n).replace("ACME", "   ACME") for n in range(5)])
    assert all("ACME" not in page for page in cleaned)


def test_stripping_is_abandoned_rather_than_gutting_a_page():
    """
    Furniture is a couple of lines. A rule that would remove half a page has
    stopped detecting headers and started deleting the document, so it bails
    entirely — noise is recoverable, missing content is not. Reachable on short
    pages, where the two edge windows overlap and every line is a candidate.
    """
    pages = [f"ACME - Internal\nRepeated body line.\nAlso repeated.\nPage {n}" for n in (1, 2, 3, 4, 5)]
    assert _strip_running_lines(pages) == pages


def test_numbered_headings_start_new_chunks():
    text = "\n".join(_section(n, f"Section {n}") for n in (1, 2, 3))
    chunks = chunk_text(text)
    assert len(chunks) == 3
    for n, chunk in zip((1, 2, 3), chunks, strict=True):
        assert chunk.content.startswith(f"{n}. Section {n}")


def test_markdown_headings_start_new_chunks():
    text = "\n".join(f"## Heading {n}\n" + ("body sentence here " * 90) for n in (1, 2, 3))
    assert len(chunk_text(text)) == 3


def test_a_section_below_the_floor_merges_forward_instead_of_standing_alone():
    # A bare heading stub embeds to almost nothing useful, so it must not become
    # a chunk of its own.
    text = "1. Tiny\nOne line.\n" + _section(2, "Substantial")
    chunks = chunk_text(text)
    assert len(chunks) == 1
    assert chunks[0].content.startswith("1. Tiny")
    assert "2. Substantial" in chunks[0].content


def test_no_chunk_falls_below_the_floor_except_the_last():
    text = "\n".join(_section(n, f"Section {n}") for n in range(1, 6))
    chunks = chunk_text(text)
    assert len(chunks) > 1
    for chunk in chunks[:-1]:
        assert chunk.token_count >= MIN_SECTION_TOKENS


def test_overlap_is_not_carried_across_a_section_boundary():
    # Overlap exists for an arbitrary mid-topic cut. A heading is the author
    # saying the topic ends, so repeating the previous section's tail would put
    # foreign text into the very embedding the split exists to sharpen.
    text = _section(1, "Alpha") + "\nUNIQUEALPHATAIL.\n" + _section(2, "Beta")
    chunks = chunk_text(text)
    assert len(chunks) == 2
    assert "UNIQUEALPHATAIL" in chunks[0].content
    assert "UNIQUEALPHATAIL" not in chunks[1].content


def test_overlap_still_applies_when_a_chunk_is_split_by_size_alone():
    # No headings here, so every boundary is a size boundary and the tail must
    # still be carried — the section rule must not disable ordinary overlap.
    text = "\n\n".join(f"Paragraph {i}. " + ("word " * 10) for i in range(30))
    chunks = chunk_text(text, max_tokens=200, overlap_tokens=60)
    assert len(chunks) > 1
    # The carried tail is whole units, so the next chunk BEGINS inside the
    # previous one — its last paragraph must reappear there.
    tail = chunks[0].content.split("\n\n")[-1]
    assert tail in chunks[1].content
    assert chunks[1].content.split("\n\n")[0] in chunks[0].content


def test_a_decimal_in_prose_is_not_mistaken_for_a_heading():
    text = "Budget rose 1. 4 percent this year.\n" + ("filler sentence text " * 120)
    assert len(chunk_text(text)) == 1


def test_control_characters_from_pdf_glyph_gaps_are_flattened():
    # pypdf hands back U+007F where a PDF has no glyph mapping — bullets and
    # en-dashes in the Afaqhims policy both arrived this way.
    chunks = chunk_text("Title \x7f Subtitle\n\x7f First bullet.\n\x7f Second bullet.")
    assert len(chunks) == 1
    assert "\x7f" not in chunks[0].content
    assert "First bullet." in chunks[0].content
